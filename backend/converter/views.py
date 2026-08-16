import textwrap
import uuid
from pathlib import Path
import math
import pdfplumber
import pypdfium2 as pdfium
from PIL import Image, ImageDraw, ImageFont
if not hasattr(Image, "ANTIALIAS"): 
    Image.ANTIALIAS = Image.Resampling.LANCZOS  
from docx import Document
from moviepy.editor import AudioFileClip, ImageClip, CompositeVideoClip, concatenate_videoclips
from rest_framework import status
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.response import Response
from rest_framework.views import APIView
from django.conf import settings

import asyncio
import edge_tts

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
MAX_FILE_SIZE_MB = 20
VIDEO_SIZE = (1280, 720)
SCROLL_PX_PER_SEC = 140 
MIN_PAGE_DURATION = 4.0
PAGE_GAP = 0
TTS_RATE = 170
TTS_VOICE = None  

def save_upload(file_obj, dest_path: Path) -> Path:
    dest_path.parent.mkdir(parents=True, exist_ok=True)
    with dest_path.open("wb+") as destination:
        for chunk in file_obj.chunks():
            destination.write(chunk)
    return dest_path

def extract_text(src_path: Path) -> str:
    suffix = src_path.suffix.lower()
    if suffix == ".txt":
        return src_path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".docx":
        doc = Document(src_path)
        return "\n".join(p.text for p in doc.paragraphs)
    if suffix == ".pdf":
        text_parts = []
        with pdfplumber.open(src_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_parts.append(page_text)
        return "\n".join(text_parts)
    raise ValueError(f"Unsupported file extension: {suffix}")

def chunk_text(text: str, max_chars: int = 320) -> list[str]:
    cleaned = " ".join(text.split())
    if not cleaned:
        return []
    return textwrap.wrap(cleaned, width=max_chars, break_long_words=False, break_on_hyphens=False)

def render_slide_image(
    text: str,
    width=VIDEO_SIZE[0],
    height=VIDEO_SIZE[1],
    margin=80,
    background=(16, 22, 37),
    font_size=32
) -> Image.Image:

    try:
        draw_font = ImageFont.truetype("arial.ttf", font_size)
    except OSError:
        draw_font = ImageFont.load_default()

    img = Image.new("RGB", (width, height), background)
    draw = ImageDraw.Draw(img)

    # Maximum text width inside the slide
    max_width = width - (margin * 2)

    # Wrap text based on actual pixel width
    lines = []

    for paragraph in text.splitlines():
        words = paragraph.split()
        current_line = ""

        for word in words:
            test_line = f"{current_line} {word}".strip()

            bbox = draw.textbbox((0, 0), test_line, font=draw_font)
            text_width = bbox[2] - bbox[0]

            if text_width <= max_width:
                current_line = test_line
            else:
                if current_line:
                    lines.append(current_line)
                current_line = word

        if current_line:
            lines.append(current_line)

    # Keep text inside the slide
    line_height = font_size + 12
    max_lines = (height - margin * 2) // line_height

    lines = lines[:max_lines]

    # Center the text vertically
    total_height = len(lines) * line_height
    y = (height - total_height) // 2

    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=draw_font)
        text_width = bbox[2] - bbox[0]

        x = (width - text_width) // 2

        draw.text(
            (x, y),
            line,
            font=draw_font,
            fill=(240, 243, 248)
        )

        y += line_height

    return img

def generate_audio(text: str, audio_path: Path):
    

    async def create_audio():
        communicate = edge_tts.Communicate(
            text,
            "en-US-AriaNeural",
            rate="+0%"
        )
        await communicate.save(str(audio_path))

    asyncio.run(create_audio())

    return AudioFileClip(str(audio_path))

def generate_audio_for_chunks(chunks: list[str], workdir: Path):
    audio_clips = []

    for index, chunk in enumerate(chunks):
        audio_path = workdir / f"audio_{index}.mp3"

        try:
            audio_clip = generate_audio(chunk, audio_path)
            audio_clips.append(audio_clip)
        except Exception as exc:
            print(f"TTS ERROR for chunk {index}: {exc}")
            audio_clips.append(None)

    return audio_clips

def render_pdf_pages(src_path: Path, workdir: Path) -> list[Path]:
    doc = pdfium.PdfDocument(str(src_path))
    image_paths: list[Path] = []
    target_width = VIDEO_SIZE[0]
    for i, page in enumerate(doc):
        scale = target_width / page.get_width()
        pil_image = page.render(scale=scale).to_pil()
        out = workdir / f"page_{i}.png"
        pil_image.save(out)
        image_paths.append(out)
    return image_paths

def render_text_doc(
    src_path: Path,
    workdir: Path,
    ext: str,
    narration_enabled
) -> tuple[list[Path], list[str]]:

    if ext == ".txt":
        text = src_path.read_text(
            encoding="utf-8",
            errors="ignore"
        )
    else:
        doc = Document(src_path)
        text = "\n".join(p.text for p in doc.paragraphs)

    chunks = chunk_text(text, max_chars=600)

    image_paths = []

    for index, chunk in enumerate(chunks):
        img = render_slide_image(chunk)

        out = workdir / f"page_{index}.png"
        img.save(out)
        img.close()

        image_paths.append(out)

    return image_paths, chunks

def build_video_from_pages(
    image_paths: list[Path],
    output_path: Path,
    audio_clips=None
):
    clips = []

    for index, img_path in enumerate(image_paths):

        # Use the matching audio duration
        if audio_clips and index < len(audio_clips) and audio_clips[index]:
            duration = audio_clips[index].duration
        else:
            duration = MIN_PAGE_DURATION

        img_clip = ImageClip(str(img_path)).set_duration(duration)

        # Attach matching voice to this slide
        if audio_clips and index < len(audio_clips) and audio_clips[index]:
            img_clip = img_clip.set_audio(audio_clips[index])

        clips.append(img_clip)

    if not clips:
        raise ValueError("No images were generated.")

    video_body = concatenate_videoclips(
        clips,
        method="chain"
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)

    video_body.write_videofile(
        str(output_path),
        fps=24,
        codec="libx264",
        audio=True,
        audio_codec="aac",
        preset="ultrafast",
        threads=1,
        verbose=False,
        logger=None,
    )

    video_body.close()

    for clip in clips:
        clip.close()

    if audio_clips:
        for audio in audio_clips:
            if audio:
                audio.close()

class HealthView(APIView):
    def get(self, _request):
        return Response({"status": "ok"})

class DocumentToVideoView(APIView):
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request):
        upload = request.FILES.get("file")
        title = request.data.get("title") or ""
        narration_enabled = request.data.get("narration") == "true"

        if not upload:
            return Response({"detail": "Attach a document file as 'file'."}, status=status.HTTP_400_BAD_REQUEST)

        if upload.size > MAX_FILE_SIZE_MB * 1024 * 1024:
            return Response({"detail": f"File too large. Limit is {MAX_FILE_SIZE_MB} MB."}, status=status.HTTP_400_BAD_REQUEST)

        ext = Path(upload.name).suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            return Response(
                {"detail": f"Unsupported format '{ext}'. Use .pdf, .docx, or .txt."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        job_id = uuid.uuid4().hex
        workdir = Path(settings.MEDIA_ROOT) / "jobs" / job_id
        src_path = workdir / f"source{ext}"
        save_upload(upload, src_path)

        try:
            if ext == ".pdf":
                image_paths = render_pdf_pages(src_path, workdir)
                chunks = []
            else:
                image_paths, chunks = render_text_doc(src_path, workdir, ext,  narration_enabled)
            video_path = workdir / "video.mp4"
            audio_clips = []

            if narration_enabled and chunks:
                try:
                    audio_clips = generate_audio_for_chunks(
                        chunks,
                        workdir
                    )
                except Exception as exc:
                    print("TTS ERROR:", exc)
                    audio_clips = []

            build_video_from_pages(image_paths, video_path, audio_clips)
        except Exception as exc:  
            return Response({"detail": f"Conversion failed: {exc}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        video_url = request.build_absolute_uri(settings.MEDIA_URL + f"jobs/{job_id}/video.mp4")
        return Response({"video_url": video_url, "job_id": job_id}, status=status.HTTP_201_CREATED)