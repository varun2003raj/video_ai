import asyncio
import subprocess
import textwrap
from pathlib import Path

import pypdfium2 as pdfium
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from moviepy.editor import AudioFileClip
import edge_tts

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
MAX_FILE_SIZE_MB = 5
MAX_PAGES = 10
VIDEO_SIZE = (854, 480)
SCROLL_PX_PER_SEC = 140
MIN_PAGE_DURATION = 4.0
PAGE_GAP = 0
TTS_RATE = 170
TTS_VOICE = None

def render_pdf_pages(src_path: Path, workdir: Path) -> list[Path]:
    doc = pdfium.PdfDocument(str(src_path))

    if len(doc) > MAX_PAGES:
        raise ValueError(
            f"PDF has too many pages. Maximum allowed is {MAX_PAGES}."
        )

    image_paths = []
    target_width = VIDEO_SIZE[0]

    for i, page in enumerate(doc):
        scale = target_width / page.get_width()
        pil_image = page.render(scale=scale).to_pil()

        out = workdir / f"page_{i}.png"
        pil_image.save(out, optimize=True)
        pil_image.close()

        image_paths.append(out)

    return image_paths


def render_text_doc(
    src_path: Path,
    workdir: Path,
    ext: str,
    narration_enabled
) -> tuple[list[Path], list[str]]:

    if ext == ".txt":
        text = src_path.read_text(
            encoding="utf-8",
            errors="ignore"
        )
    else:
        doc = Document(src_path)
        text = "\n".join(p.text for p in doc.paragraphs)

    chunks = chunk_text(text, max_chars=250)

    image_paths = []

    for index, chunk in enumerate(chunks):
        img = render_slide_image(chunk)

        out = workdir / f"page_{index}.png"
        img.save(out)
        img.close()

        image_paths.append(out)

    return image_paths, chunks


def generate_audio_for_chunks(chunks: list[str], workdir: Path):
    audio_clips = []

    for index, chunk in enumerate(chunks):
        audio_path = workdir / f"audio_{index}.mp3"

        try:
            audio_clip = generate_audio(chunk, audio_path)
            audio_clips.append(audio_clip)
        except Exception as exc:
            print(f"TTS ERROR for chunk {index}: {exc}")
            audio_clips.append(None)

    return audio_clips


def build_video_from_pages(
    image_paths: list[Path],
    output_path: Path,
    audio_clips=None
):
    if not image_paths:
        raise ValueError("No images were generated.")

    output_path.parent.mkdir(parents=True, exist_ok=True)

    temp_videos = []

    for index, img_path in enumerate(image_paths):
        if audio_clips and index < len(audio_clips) and audio_clips[index]:
            duration = max(float(audio_clips[index].duration), MIN_PAGE_DURATION)
        else:
            duration = MIN_PAGE_DURATION

        temp_video = output_path.parent / f"temp_{index}.mp4"

        # Build FFmpeg command for one slide.
        cmd = [
            "ffmpeg",
            "-y",
            "-loop", "1",
            "-i", str(img_path),
        ]

        # Add matching audio
        if audio_clips and index < len(audio_clips) and audio_clips[index]:
            audio_path = Path(audio_clips[index].filename)

            cmd.extend([
                "-i", str(audio_path),
                "-t", str(duration),
                "-map", "0:v:0",
                "-map", "1:a:0",
                "-c:v", "libx264",
                "-preset", "ultrafast",
                "-threads", "1",
                "-pix_fmt", "yuv420p",
                "-c:a", "aac",
                "-shortest",
                str(temp_video),
            ])

        else:
            cmd.extend([
                "-t", str(duration),
                "-c:v", "libx264",
                "-preset", "ultrafast",
                "-threads", "1",
                "-pix_fmt", "yuv420p",
                "-an",
                str(temp_video),
            ])

        subprocess.run(
            cmd,
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.PIPE,
            text=True,
        )

        temp_videos.append(temp_video)

    # Create FFmpeg concat file.
    concat_file = output_path.parent / "concat.txt"

    with concat_file.open("w", encoding="utf-8") as f:
        for video in temp_videos:
            f.write(f"file '{video.as_posix()}'\n")

    # Join the already encoded clips without re-encoding.
    subprocess.run(
        [
            "ffmpeg",
            "-y",
            "-f", "concat",
            "-safe", "0",
            "-i", str(concat_file),
            "-c:v", "libx264",
            "-preset", "ultrafast",
            "-threads", "1",
            "-pix_fmt", "yuv420p",
            "-c:a", "aac",
            "-b:a", "128k",
            "-movflags", "+faststart",
            str(output_path),
        ],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.PIPE,
        text=True,
    )

    # Clean temporary videos.
    for video in temp_videos:
        try:
            video.unlink()
        except OSError:
            pass

    try:
        concat_file.unlink()
    except OSError:
        pass


def chunk_text(text: str, max_chars: int = 320) -> list[str]:
    cleaned = " ".join(text.split())
    if not cleaned:
        return []
    return textwrap.wrap(cleaned, width=max_chars, break_long_words=False, break_on_hyphens=False)

def extract_text(src_path: Path) -> str:
    suffix = src_path.suffix.lower()

    if suffix == ".txt":
        return src_path.read_text(
            encoding="utf-8",
            errors="ignore"
        )

    if suffix == ".docx":
        doc = Document(src_path)
        return "\n".join(
            p.text for p in doc.paragraphs
        )

    if suffix == ".pdf":
        import pdfplumber

        text_parts = []

        with pdfplumber.open(src_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_parts.append(page_text)

        return "\n".join(text_parts)

    raise ValueError(
        f"Unsupported file extension: {suffix}"
    )

def render_slide_image(
    text: str,
    width=VIDEO_SIZE[0],
    height=VIDEO_SIZE[1],
    margin=80,
    background=(16, 22, 37),
    font_size=32
) -> Image.Image:

    try:
        draw_font = ImageFont.truetype("arial.ttf", font_size)
    except OSError:
        draw_font = ImageFont.load_default()

    img = Image.new("RGB", (width, height), background)
    draw = ImageDraw.Draw(img)

    # Maximum text width inside the slide
    max_width = width - (margin * 2)

    # Wrap text based on actual pixel width
    lines = []

    for paragraph in text.splitlines():
        words = paragraph.split()
        current_line = ""

        for word in words:
            test_line = f"{current_line} {word}".strip()

            bbox = draw.textbbox((0, 0), test_line, font=draw_font)
            text_width = bbox[2] - bbox[0]

            if text_width <= max_width:
                current_line = test_line
            else:
                if current_line:
                    lines.append(current_line)
                current_line = word

        if current_line:
            lines.append(current_line)

    # Keep text inside the slide
    line_height = font_size + 12
    max_lines = (height - margin * 2) // line_height

    

    # Center the text vertically
    total_height = len(lines) * line_height
    y = (height - total_height) // 2

    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=draw_font)
        text_width = bbox[2] - bbox[0]

        x = (width - text_width) // 2

        draw.text(
            (x, y),
            line,
            font=draw_font,
            fill=(240, 243, 248)
        )

        y += line_height

    return img



def generate_audio(text: str, audio_path: Path):
    

    async def create_audio():
        communicate = edge_tts.Communicate(
            text,
            "en-US-AriaNeural",
            rate="+0%"
        )
        await communicate.save(str(audio_path))

    asyncio.run(create_audio())

    return AudioFileClip(str(audio_path))